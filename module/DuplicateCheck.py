# 导入所需的模块和库
import os
import nltk
import numpy as np
from docx import Document
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from PyQt5.QtWidgets import QProgressBar

# 第一次运行的时候才用，下载不了的话翻墙可能是服务器在外网，挂梯子就行
# nltk.download('punkt')  # 下载nltk的分词模型

# 文档查重类
class Duplication:
    def __init__(self, new_document_path, library_path):
        self.new_document_path = new_document_path
        self.library_path = library_path

    # 从段落中获取句子
    def __get_sentences_from_paragraph(self, paragraph):
        return nltk.sent_tokenize(paragraph.text)

    # 从文档中获取文本内容
    def __get_text_content_from_document(self, doc):
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.runs and any(run.text.strip() for run in paragraph.runs):
                sentences = self.__get_sentences_from_paragraph(paragraph)
                content.extend(sentences)
        return content

    # 计算余弦相似度
    def __calculate_cosine_similarity(self, original_content, new_sentence):
        vectorizer = CountVectorizer().fit_transform(original_content + [new_sentence])
        vectors = vectorizer.toarray()
        original_vectors = vectors[:-1]
        new_vector = vectors[-1]

        similarity_scores = cosine_similarity(original_vectors, [new_vector])
        return similarity_scores

    # 寻找最相似的句子
    def __find_most_similar_sentence(self, original_content, new_content, similarity_threshold=0.7):
        most_similar_sentence = None
        most_similar_original_sentence = None
        max_similarity = 0

        for new_sentence in new_content:
            similarity_scores = self.__calculate_cosine_similarity(original_content, new_sentence)
            most_similar_score = np.max(similarity_scores)

            if most_similar_score > similarity_threshold and most_similar_score > max_similarity:
                max_similarity = most_similar_score
                most_similar_index = np.argmax(similarity_scores)
                most_similar_original_sentence = original_content[most_similar_index]
                most_similar_sentence = new_sentence

        return most_similar_sentence, most_similar_original_sentence, round(max_similarity, 2)

    # 执行查重操作
    def work(self, progressBar: QProgressBar):
        # 原始文档库路径
        library_path = self.library_path

        # 读取原始文档库
        original_documents = []
        document_names = []  # 保存文档名
        for filename in os.listdir(library_path):
            if filename.endswith(".docx") and not filename.startswith("~$"):
                document_path = os.path.join(library_path, filename)
                original_doc = Document(document_path)
                original_content = self.__get_text_content_from_document(original_doc)
                original_documents.extend(original_content)
                document_names.extend([filename] * len(original_content))

        # 读取新文档
        new_document_path = self.new_document_path
        new_doc = Document(new_document_path)
        new_content = self.__get_text_content_from_document(new_doc)
        similar_content_length = 0
        return_list = []

        # 与每个原始文档比较
        progressBar.setRange(0, len(new_content))
        for i, new_sentence in enumerate(new_content):
            most_similar_sentence, most_similar_original_sentence, max_similarity = self.__find_most_similar_sentence(
                original_documents, [new_sentence])

            # 输出结果
            if most_similar_sentence:
                similar_document_name = document_names[original_documents.index(most_similar_original_sentence)]
                print(
                    f"\n新文档句子：'{most_similar_sentence}'\n最相似的句子：'{most_similar_original_sentence}'\n最大相似度：{max_similarity}\n所在文档：{similar_document_name}")
                similar_content_length += len(most_similar_sentence) * max_similarity
                return_list.append(
                    (most_similar_sentence, most_similar_original_sentence, max_similarity, similar_document_name))
            else:
                similar_content_length += len(new_sentence) * 0
            progressBar.setValue(i + 1)

        total_length = 0
        for paragraph in new_doc.paragraphs:
            total_length += len(paragraph.text)
        Duplication_rate = similar_content_length / total_length
        print("查重率为：", Duplication_rate)
        return return_list, Duplication_rate

# 创建Duplication类的实例
def create_checker(new_document_path, library_path):
    return Duplication(new_document_path, library_path)

if __name__ == "__main__":
    # 从环境变量中获取文档库路径和新文档路径
    library_path = os.environ['document_library_path']
    new_document_path = os.environ['new_document_path']

    # 创建Duplication类的实例并调用work方法
    dp = Duplication(new_document_path, library_path)
    dp.work()
